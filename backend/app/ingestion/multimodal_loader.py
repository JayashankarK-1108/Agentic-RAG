
import os
import zipfile
import fitz
from docx import Document
from docx.oxml.ns import qn
from app.ingestion.s3_uploader import upload_image


# ── PDF extraction ─────────────────────────────────────────

def extract_pdf(file_path):
    """
    Extract text and images from a PDF file.
    Each page becomes one step; images on that page are attached to it.

    Returns:
        pages      : list of {text, images}
        image_stats: {found, uploaded, failed}
    """
    try:
        source = os.path.basename(file_path)
        doc = fitz.open(file_path)
        pages = []
        found = uploaded = failed = 0

        for pno, page in enumerate(doc):
            text = page.get_text()
            imgs = []
            for i, img in enumerate(page.get_images(full=True)):
                found += 1
                try:
                    xref = img[0]
                    base = doc.extract_image(xref)
                    url = upload_image(base["image"], f"page{pno}_{i}.png", source=source)
                    if url:
                        imgs.append(url)
                        uploaded += 1
                    else:
                        failed += 1
                except Exception as e:
                    failed += 1
                    print(f"Error extracting image from page {pno}: {e}")
            pages.append({"text": text, "images": imgs})

        return pages, {"found": found, "uploaded": uploaded, "failed": failed}
    except Exception as e:
        print(f"Error extracting PDF: {e}")
        return [], {"found": 0, "uploaded": 0, "failed": 0}


# ── DOCX extraction ────────────────────────────────────────

def _upload_docx_images(file_path):
    """
    Pre-upload every image inside the DOCX zip (word/media/*) to S3.
    Returns a dict mapping relationship ID → S3 URL, and image stats.
    """
    rel_to_url = {}
    found = uploaded = failed = 0
    source = os.path.basename(file_path)

    try:
        doc = Document(file_path)
        with zipfile.ZipFile(file_path) as zf:
            media = {os.path.basename(n): zf.read(n)
                     for n in zf.namelist() if n.startswith("word/media/")}

        for rid, rel in doc.part.rels.items():
            if "image" not in rel.reltype:
                continue
            img_name = os.path.basename(rel.target_ref)
            img_data = media.get(img_name)
            if img_data is None:
                continue
            found += 1
            url = upload_image(img_data, img_name, source=source)
            if url:
                rel_to_url[rid] = url
                uploaded += 1
            else:
                failed += 1

    except Exception as e:
        print(f"Error uploading DOCX images: {e}")

    return rel_to_url, {"found": found, "uploaded": uploaded, "failed": failed}


def _get_paragraph_image_rids(para):
    """Return relationship IDs for all inline images inside a paragraph."""
    rids = []
    for blip in para._element.iter(qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if rid:
            rids.append(rid)
    return rids


def extract_docx(file_path):
    """
    Extract text and images from a DOCX file.

    Paragraphs are processed in document order. Each step is a block of
    text paragraphs followed by their image(s). When an image paragraph
    is encountered, the accumulated text + that image are saved as one
    step and we start fresh — so each image stays with the step it
    belongs to in the document.

    Returns:
        pages      : list of {text, images}
        image_stats: {found, uploaded, failed}
    """
    try:
        rel_to_url, img_stats = _upload_docx_images(file_path)
        doc = Document(file_path)

        pages = []
        pending_texts = []   # text lines accumulated since last image
        pending_images = []  # images accumulated for current text block

        def flush(extra_images=None):
            """Save current pending text + images as one step."""
            imgs = list(pending_images) + (extra_images or [])
            if pending_texts:
                pages.append({
                    "text":   "\n".join(pending_texts),
                    "images": imgs,
                })
                pending_texts.clear()
                pending_images.clear()

        for para in doc.paragraphs:
            text = para.text.strip()
            para_images = [
                rel_to_url[rid]
                for rid in _get_paragraph_image_rids(para)
                if rid in rel_to_url
            ]

            if para_images:
                # This paragraph contains an image.
                # If it also has text, include that text in the current block.
                if text:
                    pending_texts.append(text)
                # Save the accumulated text with this image, then reset.
                flush(extra_images=para_images)
            elif text:
                # Text-only paragraph — accumulate.
                pending_texts.append(text)

        # Flush any trailing text that has no following image
        flush()

        # Also extract table rows (no specific image association)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    pages.append({"text": row_text, "images": []})

        return pages, img_stats

    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        return [], {"found": 0, "uploaded": 0, "failed": 0}


# ── Universal entry point ──────────────────────────────────

def extract_document(file_path):
    """Auto-detect file type and extract text + images."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_docx(file_path)
    else:
        print(f"Unsupported file type: {ext}")
        return [], {"found": 0, "uploaded": 0, "failed": 0}
